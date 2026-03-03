"""DOCX format plugin using python-docx."""

from pathlib import Path
from typing import Optional

from ion.core.exceptions import PluginError


class DocxPlugin:
    """Plugin for DOCX (Word) format."""

    @property
    def name(self) -> str:
        return "docx"

    @property
    def extensions(self) -> list[str]:
        return [".docx"]

    def read(self, path: Path) -> str:
        """Read content from a DOCX file as plain text."""
        try:
            from docx import Document
        except ImportError:
            raise PluginError("docx", "python-docx is required for DOCX support")

        doc = Document(str(path))
        paragraphs = []

        for para in doc.paragraphs:
            paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    def write(self, content: str, path: Path) -> None:
        """Write content to a DOCX file."""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise PluginError("docx", "python-docx is required for DOCX support")

        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Parse content and create appropriate elements
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Handle headers (markdown-style)
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.strip() == "":
                pass  # Skip empty lines
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

            i += 1

        doc.save(str(path))

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with basic markdown formatting to a paragraph."""
        import re

        # Pattern for bold, italic, and code
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
            else:
                paragraph.add_run(part)

    def convert(self, content: str, from_format: str) -> str:
        """Convert content from another format.

        Note: For DOCX, we return the content as-is since writing
        handles the conversion. The actual DOCX creation happens in write().
        """
        return content

    def read_structured(self, path: Path) -> dict:
        """Read DOCX with structure information preserved."""
        try:
            from docx import Document
        except ImportError:
            raise PluginError("docx", "python-docx is required for DOCX support")

        doc = Document(str(path))
        result = {
            "paragraphs": [],
            "tables": [],
            "headers": [],
        }

        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name if para.style else None,
            }
            result["paragraphs"].append(para_info)

            # Track headers
            if para.style and "Heading" in para.style.name:
                result["headers"].append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append(table_data)

        return result
