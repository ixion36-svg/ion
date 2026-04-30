"""Translator service (v0.17.0).

Two responsibilities:

1. **Text translation** via the configured Ollama model. Default is
   ``ION_TRANSLATOR_MODEL`` falling back to ``ION_OLLAMA_MODEL`` falling
   back to ``llama3.1:8b``. Llama 3.1 + qwen2.5 are both multilingual
   enough for typical SOC text (decoded payloads, phishing content,
   threat-actor messages, intel reports).

2. **Text extraction from uploaded documents** — txt, md, csv, log, json,
   html, eml, docx, pdf. Returns plain text the caller can pass into
   ``translate_text``.

Design principles:

- **Air-gapped first**: every dep is either stdlib or already in
  pyproject.toml. ``pypdf`` is the only addition.
- **Sync API**: callers (FastAPI handlers) wrap in ``asyncio.to_thread``
  if they need to run alongside other async work.
- **Hard input cap**: ``MAX_INPUT_CHARS`` to keep us inside the LLM's
  context. Long docs are chunked and concatenated by the API layer.
- **Graceful degradation**: extraction failures return ``""`` with a
  log line; translation failures raise ``TranslationError`` so the
  HTTP route can return 502.
"""

from __future__ import annotations

import io
import logging
import os
import re
from typing import Dict, List, Optional, Tuple

import httpx

logger = logging.getLogger(__name__)


# ── Config knobs ─────────────────────────────────────────────────────────
MAX_INPUT_CHARS = 10_000     # per LLM call; longer docs chunk in the API
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB upload cap
LLM_TIMEOUT_SECONDS = 120.0  # llama3.1:8b on CPU can be slow on long input


# Languages we expose in the picker. ISO 639-1 + a friendly label. Order
# is the order the picker shows them — most-likely SOC sources first.
SUPPORTED_LANGUAGES: List[Dict[str, str]] = [
    {"code": "auto", "label": "Auto-detect"},
    {"code": "en",   "label": "English"},
    {"code": "ru",   "label": "Russian"},
    {"code": "zh",   "label": "Chinese (Simplified)"},
    {"code": "ja",   "label": "Japanese"},
    {"code": "ko",   "label": "Korean"},
    {"code": "ar",   "label": "Arabic"},
    {"code": "fa",   "label": "Persian / Farsi"},
    {"code": "es",   "label": "Spanish"},
    {"code": "fr",   "label": "French"},
    {"code": "de",   "label": "German"},
    {"code": "pt",   "label": "Portuguese"},
    {"code": "it",   "label": "Italian"},
    {"code": "tr",   "label": "Turkish"},
    {"code": "vi",   "label": "Vietnamese"},
]


_LANG_BY_CODE = {l["code"]: l["label"] for l in SUPPORTED_LANGUAGES}


class TranslationError(Exception):
    """Raised on Ollama failure — the HTTP layer maps to 502."""


# ── Provider config ──────────────────────────────────────────────────────


def _ollama_url() -> str:
    return (
        os.environ.get("ION_OLLAMA_URL")
        or "http://ion-ollama:11434"
    ).rstrip("/")


def _translator_model() -> str:
    return (
        os.environ.get("ION_TRANSLATOR_MODEL")
        or os.environ.get("ION_OLLAMA_MODEL")
        or "llama3.1:8b"
    )


# ── LLM-backed translation ───────────────────────────────────────────────


def _build_prompt(text: str, source_lang: str, target_lang: str) -> str:
    """Build a translation prompt that nudges the LLM to return ONLY
    the translation — no preamble, no commentary, no quotes."""
    src_label = (
        "auto-detect"
        if source_lang == "auto"
        else _LANG_BY_CODE.get(source_lang, source_lang)
    )
    tgt_label = _LANG_BY_CODE.get(target_lang, target_lang)
    return (
        f"You are a professional translator. Translate the following text "
        f"from {src_label} into {tgt_label}.\n\n"
        f"Return ONLY the translation. Do not add any preamble, explanation, "
        f"quotes, or commentary. Preserve technical terms (CVE IDs, hashes, "
        f"IP addresses, file paths, command-line flags) verbatim. Preserve "
        f"line breaks where present.\n\n"
        f"---\n{text}\n---"
    )


def translate_text(
    text: str,
    target_lang: str = "en",
    source_lang: str = "auto",
    *,
    timeout: Optional[float] = None,
) -> str:
    """Translate ``text`` to ``target_lang``. Returns the translated string.

    Raises ``TranslationError`` on any failure — the caller (HTTP layer)
    is expected to surface a meaningful error to the user. Empty input
    returns ``""`` without calling the LLM.
    """
    if not text or not text.strip():
        return ""
    if target_lang not in _LANG_BY_CODE or target_lang == "auto":
        raise TranslationError(f"Invalid target language: {target_lang!r}")
    if source_lang not in _LANG_BY_CODE:
        raise TranslationError(f"Invalid source language: {source_lang!r}")
    if len(text) > MAX_INPUT_CHARS:
        # The API layer should chunk before calling — defensive truncate.
        text = text[:MAX_INPUT_CHARS]

    prompt = _build_prompt(text, source_lang, target_lang)
    body = {
        "model": _translator_model(),
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.2,    # tight — translation isn't creative
            "num_predict": 4096,   # roughly 1 output token per source token
        },
    }
    try:
        with httpx.Client(timeout=timeout or LLM_TIMEOUT_SECONDS) as client:
            r = client.post(f"{_ollama_url()}/api/generate", json=body)
            r.raise_for_status()
            data = r.json()
    except (httpx.HTTPError, ValueError) as exc:
        logger.warning("Translator LLM call failed: %s", exc)
        raise TranslationError(f"LLM call failed: {exc}") from exc

    out = (data.get("response") or "").strip()
    # The model occasionally wraps the output in --- fences mirroring the
    # prompt; strip those so the user sees clean output.
    out = re.sub(r"^---+\s*\n?", "", out)
    out = re.sub(r"\n?---+\s*$", "", out)
    return out.strip()


# ── Document text extraction ─────────────────────────────────────────────


# (extension, kind) — `kind` selects the extractor below
_EXT_KIND = {
    ".txt":  "plain",
    ".md":   "plain",
    ".csv":  "plain",
    ".log":  "plain",
    ".json": "plain",
    ".html": "html",
    ".htm":  "html",
    ".eml":  "eml",
    ".msg":  "eml",   # best-effort; outlook .msg falls back to plain text
    ".docx": "docx",
    ".pdf":  "pdf",
}


def _decode_text(content: bytes) -> str:
    """Decode bytes as UTF-8 with utf-8-sig BOM tolerance, falling back
    to latin-1 (always succeeds) if utf-8 fails. Preserves the original
    line breaks so structure isn't lost."""
    if content.startswith(b"\xef\xbb\xbf"):
        content = content[3:]
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("utf-16")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="replace")


def _extract_plain(content: bytes) -> str:
    return _decode_text(content)


def _extract_html(content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
    except Exception as exc:
        logger.warning("BeautifulSoup unavailable: %s", exc)
        return _decode_text(content)
    soup = BeautifulSoup(_decode_text(content), "html.parser")
    # Drop script/style — they're noise for translation.
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    # Collapse runs of blank lines so the LLM doesn't waste tokens on them.
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _extract_eml(content: bytes) -> str:
    """Pull subject + from + to + plain-text body from an RFC 5322 message."""
    from email import policy
    from email.parser import BytesParser
    msg = BytesParser(policy=policy.default).parsebytes(content)
    parts: List[str] = []
    subject = msg.get("subject") or ""
    sender = msg.get("from") or ""
    recipient = msg.get("to") or ""
    if subject:
        parts.append(f"Subject: {subject}")
    if sender:
        parts.append(f"From: {sender}")
    if recipient:
        parts.append(f"To: {recipient}")
    parts.append("")  # blank line
    # Prefer text/plain body part; fall back to text/html if that's all there is
    plain = msg.get_body(preferencelist=("plain",))
    if plain is not None:
        try:
            parts.append(plain.get_content())
        except Exception:
            pass
    else:
        html_part = msg.get_body(preferencelist=("html",))
        if html_part is not None:
            try:
                parts.append(_extract_html(html_part.get_content().encode("utf-8")))
            except Exception:
                pass
    return "\n".join(parts).strip()


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        logger.warning("python-docx unavailable: %s", exc)
        raise
    bio = io.BytesIO(content)
    doc = Document(bio)
    paras = [p.text for p in doc.paragraphs if p.text]
    # Tables can contain alert text in some templates — pull cell text too.
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras).strip()


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        logger.warning("pypdf unavailable — install `pypdf>=4.0`: %s", exc)
        raise
    bio = io.BytesIO(content)
    reader = PdfReader(bio)
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            logger.debug("pypdf page extract failed: %s", exc)
            pages.append("")
    return "\n\n".join(p.strip() for p in pages if p.strip()).strip()


_EXTRACTORS = {
    "plain": _extract_plain,
    "html":  _extract_html,
    "eml":   _extract_eml,
    "docx":  _extract_docx,
    "pdf":   _extract_pdf,
}


def supported_extensions() -> List[str]:
    return sorted(_EXT_KIND.keys())


def extract_text_from_upload(filename: str, content: bytes) -> Tuple[str, str]:
    """Extract translatable text from an uploaded file.

    Returns ``(text, kind)`` where ``kind`` is one of ``plain | html |
    eml | docx | pdf``. Raises ``ValueError`` for unsupported extensions
    and lets the underlying extractor's exception propagate on parse
    errors so the API can surface a useful error message.
    """
    if not filename:
        raise ValueError("filename is required")
    if len(content) > MAX_FILE_BYTES:
        raise ValueError(
            f"File too large: {len(content)} bytes (max {MAX_FILE_BYTES})"
        )
    ext = ""
    lower = filename.lower()
    for e in _EXT_KIND:
        if lower.endswith(e):
            ext = e
            break
    if not ext:
        raise ValueError(
            f"Unsupported file type. Allowed: {', '.join(sorted(_EXT_KIND))}"
        )
    kind = _EXT_KIND[ext]
    extractor = _EXTRACTORS[kind]
    text = extractor(content)
    return text, kind


# ── Long-document chunking ───────────────────────────────────────────────


def chunk_text(text: str, max_chars: int = MAX_INPUT_CHARS) -> List[str]:
    """Split a long document into chunks ≤ ``max_chars``, preferring
    paragraph then sentence boundaries so the LLM doesn't translate a
    half-sentence on either side of a chunk break."""
    if len(text) <= max_chars:
        return [text]
    chunks: List[str] = []
    paragraphs = re.split(r"\n\s*\n", text)
    current = ""
    for p in paragraphs:
        if len(current) + len(p) + 2 <= max_chars:
            current = (current + "\n\n" + p) if current else p
            continue
        if current:
            chunks.append(current)
            current = ""
        # Paragraph itself might be > max_chars — split on sentence ends.
        if len(p) > max_chars:
            sentences = re.split(r"(?<=[.!?])\s+", p)
            buf = ""
            for s in sentences:
                if len(buf) + len(s) + 1 <= max_chars:
                    buf = (buf + " " + s).strip()
                else:
                    if buf:
                        chunks.append(buf)
                    if len(s) > max_chars:
                        # Hard wrap at max_chars as a last resort.
                        for i in range(0, len(s), max_chars):
                            chunks.append(s[i : i + max_chars])
                        buf = ""
                    else:
                        buf = s
            if buf:
                current = buf
        else:
            current = p
    if current:
        chunks.append(current)
    return chunks


def translate_long_text(
    text: str,
    target_lang: str = "en",
    source_lang: str = "auto",
) -> str:
    """Translate text that may exceed ``MAX_INPUT_CHARS`` by chunking on
    paragraph/sentence boundaries and re-joining the translations."""
    chunks = chunk_text(text)
    if len(chunks) == 1:
        return translate_text(chunks[0], target_lang, source_lang)
    out_parts: List[str] = []
    for i, chunk in enumerate(chunks, 1):
        logger.debug("Translator: chunk %d/%d (%d chars)", i, len(chunks), len(chunk))
        out_parts.append(translate_text(chunk, target_lang, source_lang))
    return "\n\n".join(out_parts)
